import json
from Pynite import FEModel3D
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np
import os

# --- Helper Functions ---
def parse_support_conditions(support_str):
    """
    Parses the support condition string (e.g., 'xyz,xyz' or 'xyz,y' or 'XYZ,xyz')
    into PyniteFEA compatible format.

    Lowercase 'x', 'y', 'z' define translational restraints (True if present).
    Uppercase 'X', 'Y', 'Z' define rotational restraints (True if present).

    Examples:
    - "xyz": Pinned support (UX, UY, UZ = True; RX, RY, RZ = False)
    - "y": Roller support allowing X, Z translation and all rotations (UY = True)
    - "XYZ": Fixed support for rotations only (RX, RY, RZ = True; UX, UY, UZ = False - unusual standalone)
    - "xyzXYZ": Fully fixed support (all True)

    PyniteFEA def_support arguments: UX, UY, UZ, RX, RY, RZ (True for fixed, False for released)
    """
    conditions = []
    parts = support_str.split(',')
    for part in parts:
        fixity = {
            'UX': 'x' in part, 'UY': 'y' in part, 'UZ': 'z' in part,
            'RX': 'X' in part, 'RY': 'Y' in part, 'RZ': 'Z' in part
        }
        conditions.append(fixity)
    return conditions

def generate_nodes(span, segment_length=0.1):
    """Generates node coordinates: support nodes (N0, N1) and internal segment nodes."""
    nodes = {'N0': (0, 0, 0), 'N1': (span, 0, 0)}
    internal_nodes = {}
    num_segments = int(span / segment_length)
    for i in range(num_segments + 1): # Include the start and end points of segments
        x_coord = i * segment_length
        # Ensure we don't duplicate N0 or N1 if segment length aligns perfectly
        if not np.isclose(x_coord, 0) and not np.isclose(x_coord, span):
             # Check if x_coord is too close to span before adding
            if x_coord < span:
                 internal_nodes[f'n{i}'] = (x_coord, 0, 0)

    # Sort internal nodes by x-coordinate to ensure order for plotting
    sorted_internal_node_names = sorted(internal_nodes.keys(), key=lambda name: internal_nodes[name][0])

    # Combine all nodes: N0, sorted internal nodes, N1
    all_node_coords = {'N0': nodes['N0']}
    for name in sorted_internal_node_names:
        all_node_coords[name] = internal_nodes[name]
    all_node_coords['N1'] = nodes['N1']

    return all_node_coords

def add_load_combination_factors(doc, combo_type, combinations_data):
    """Adds load combination factors to the DOCX document."""
    doc.add_heading(f'{combo_type} Load Combinations', level=2)
    # No. of combinations: {combinations_data[f'No. of {combo_type} combinations']}"

    # Table for combinations
    num_combos = combinations_data[f'No. of {combo_type} combinations']
    combos = combinations_data['Combinations']

    if not combos:
        doc.add_paragraph("No combinations defined.")
        return

    # Determine headers from the keys of the first combination, assuming consistency
    # Ordered headers for better readability
    factor_keys_ordered = ["Permanent factor", "Live factor", "Snow factor", "Wind factor"]
    headers = ["Combination"] + factor_keys_ordered

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header_name in enumerate(headers):
        hdr_cells[i].text = header_name

    for i, combo in enumerate(combos):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{combo_type}{i+1}"
        for j, key in enumerate(factor_keys_ordered):
            row_cells[j+1].text = str(combo.get(key, 0)) # Use .get for safety

def plot_beam_geometry_and_loads(beam_props, beam_loads, output_path='plots/beam_geometry_loads.png'):
    """Plots the beam geometry, supports, and applied loads using matplotlib."""
    span = float(beam_props['Span'])
    support_conditions_str = beam_props['Support conditions']

    fig, ax = plt.subplots(figsize=(12, 4)) # Adjusted figure size for better layout

    # Beam line
    ax.plot([0, span], [0, 0], 'k-', lw=3) # Beam as a thick black line at y=0

    # --- Define Y-offsets and symbol sizes relative to span for responsiveness ---
    base_y_offset = 0.05 * span  # Base offset for symbols from beam line
    text_y_offset = base_y_offset * 1.3 # Offset for text below symbols
    load_arrow_head_width = 0.02 * span
    load_arrow_head_length = 0.03 * span
    udl_rect_height = 0.06 * span # Height of UDL rectangle above beam
    point_load_arrow_length = udl_rect_height * 1.5 # Length of point load arrow stem

    # Support symbols
    support_parts = support_conditions_str.split(',')

    # Support N0 (at x=0)
    if len(support_parts) > 0:
        s_type_raw = support_parts[0] # Keep raw for text label
        s_type_proc = s_type_raw.strip() # Processed for logic, no .lower() to distinguish X/Y/Z from x/y/z

        has_x_trans = 'x' in s_type_proc
        has_y_trans = 'y' in s_type_proc
        has_z_trans = 'z' in s_type_proc
        has_any_trans = has_x_trans or has_y_trans or has_z_trans

        has_x_rot = 'X' in s_type_proc
        has_y_rot = 'Y' in s_type_proc
        has_z_rot = 'Z' in s_type_proc
        has_any_rot = has_x_rot or has_y_rot or has_z_rot

        if has_any_rot: # If any rotational restraint, draw as Fixed
            ax.plot([0,0], [-base_y_offset*0.5, base_y_offset*0.5], 'k-', lw=2)
            ax.fill_between([-0.05*span, 0], [-base_y_offset*0.5, -base_y_offset*0.5], [base_y_offset*0.5, base_y_offset*0.5], color='dimgray', hatch='///')
        elif has_y_trans and not has_x_trans and not has_z_trans: # Roller Y (only Y translation fixed, no rotation)
            ax.plot(0, 0, 'ko', markersize=10, mfc='white')
            ax.plot([-0.02*span, 0.02*span], [-base_y_offset*0.3, -base_y_offset*0.3], 'k-', lw=1) # Corrected Y for line
        elif has_any_trans: # Pinned (any translational fixed, no rotation)
            ax.plot(0, 0, 'k^', markersize=10)
        else: # Default/Other or no restraint
            ax.plot(0, 0, 'ks', markersize=8) # Square for other/custom or no symbol if preferred
        ax.text(0, -text_y_offset, s_type_raw, ha='center', va='top')


    # Support N1 (at x=span)
    if len(support_parts) > 1:
        s_type_raw = support_parts[1] # Keep raw for text label
        s_type_proc = s_type_raw.strip()

        has_x_trans = 'x' in s_type_proc
        has_y_trans = 'y' in s_type_proc
        has_z_trans = 'z' in s_type_proc
        has_any_trans = has_x_trans or has_y_trans or has_z_trans

        has_x_rot = 'X' in s_type_proc
        has_y_rot = 'Y' in s_type_proc
        has_z_rot = 'Z' in s_type_proc
        has_any_rot = has_x_rot or has_y_rot or has_z_rot

        if has_any_rot: # If any rotational restraint, draw as Fixed
            ax.plot([span,span], [-base_y_offset*0.5, base_y_offset*0.5], 'k-', lw=2)
            ax.fill_between([span, span+0.05*span], [-base_y_offset*0.5, -base_y_offset*0.5], [base_y_offset*0.5, base_y_offset*0.5], color='dimgray', hatch='///')
        elif has_y_trans and not has_x_trans and not has_z_trans: # Roller Y
            ax.plot(span, 0, 'ko', markersize=10, mfc='white')
            ax.plot([span - 0.02*span, span + 0.02*span], [-base_y_offset*0.3, -base_y_offset*0.3], 'k-', lw=1) # Corrected Y for line
        elif has_any_trans: # Pinned
            ax.plot(span, 0, 'k^', markersize=10)
        else: # Default/Other
            ax.plot(span, 0, 'ks', markersize=8)
        ax.text(span, -text_y_offset, s_type_raw, ha='center', va='top')

    # Applied UDLs
    for udl in beam_loads.get('UDL', []): # Use .get for robustness
        mag = udl['Magnitude']
        load_name = udl['Load name']
        start_x = udl['Start'] if udl['Full/Partial'].lower() == 'partial' and udl['Start'] is not None else 0
        end_x = udl['End'] if udl['Full/Partial'].lower() == 'partial' and udl['End'] is not None else span # Corrected indentation

        # Draw rectangle for UDL (above the beam line)
        ax.add_patch(plt.Rectangle((start_x, 0), end_x - start_x, udl_rect_height, facecolor='skyblue', edgecolor='dodgerblue', alpha=0.7))

        # Arrows for UDL representation (pointing downwards onto the rectangle)
        num_udl_arrows = max(2, int((end_x - start_x) / (span/8))) # Adjust arrow density
        for i in range(num_udl_arrows + 1):
            arrow_x = start_x + i * (end_x - start_x) / num_udl_arrows
            # Arrow from slightly above rectangle, pointing into it
            ax.arrow(arrow_x, udl_rect_height * 1.1, 0, -udl_rect_height * 0.3,
                     head_width=load_arrow_head_width*0.7, head_length=load_arrow_head_length*0.7,
                     fc='dodgerblue', ec='dodgerblue', lw=0.8)

        ax.text((start_x + end_x) / 2, udl_rect_height * 1.2, f"{load_name}: {mag} kN/m",
                ha='center', va='bottom', color='dodgerblue', fontsize=9)

    # Applied Point Loads
    for pl_idx, pl in enumerate(beam_loads.get('Point Load', [])): # Use .get for robustness
        mag = pl['Magnitude']
        load_name = pl['Load name']
        pos_x = pl['Start']
        # Stagger text for closely spaced point loads if needed (simple y-offset here)
        text_y_pl_offset = point_load_arrow_length * 1.1 + (pl_idx % 2 * base_y_offset * 0.8)

        ax.arrow(pos_x, point_load_arrow_length, 0, -point_load_arrow_length,
                     head_width=load_arrow_head_width, head_length=load_arrow_head_length,
                     fc='crimson', ec='crimson', lw=1.5)
        ax.text(pos_x, text_y_pl_offset, f"{load_name}: {mag} kN",
                ha='center', va='bottom', color='crimson', fontsize=9) # Corrected indentation

    ax.set_xlim(-0.1 * span, 1.1 * span)
    # Adjust ylim to ensure all elements are visible
    min_y_lim = -text_y_offset - base_y_offset # Space for support text
    max_y_lim = udl_rect_height * 1.5 + base_y_offset # Space for UDL text and point loads
    if 'Point Load' in beam_loads and beam_loads['Point Load']: # Check if point loads exist
         max_y_lim = max(max_y_lim, point_load_arrow_length * 1.2 + max( (idx % 2 * base_y_offset*0.8) for idx in range(len(beam_loads['Point Load'])) ) + base_y_offset)


    ax.set_ylim(min_y_lim, max_y_lim)
    ax.set_xlabel("Position along beam (m)")
    ax.set_yticks([]) # Hide y-axis ticks as it's schematic
    ax.set_title("Beam Configuration: Geometry, Supports, and Loads", fontsize=12)
    ax.spines['left'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.tight_layout(pad=1.5) # Add padding

    # Ensure plots directory exists
    plots_dir = os.path.dirname(output_path)
    if not os.path.exists(plots_dir) and plots_dir: # Check if plots_dir is not empty
        os.makedirs(plots_dir)

    plt.savefig(output_path)
    plt.close(fig)
    print(f"Beam geometry diagram saved to {output_path}")


# --- Main Analysis Function ---
def run_beam_analysis(input_file='beam_input.json', output_docx='beam_analysis_report.docx'):
    """Runs the entire beam analysis and generates a DOCX report."""

    # Create a directory for plots if it doesn't exist
    if not os.path.exists('plots'):
        os.makedirs('plots')

    # 1. Load Input Data
    with open(input_file, 'r') as f:
        data = json.load(f)

    beam_props = data['Beam Properties']
    timber_section = data['Timber Section']
    beam_loads = data['Beam Loads']
    load_combos = data['Load combinations']

    span = float(beam_props['Span'])

    # ** Plot beam geometry and loads **
    plot_beam_geometry_and_loads(beam_props, beam_loads, output_path='plots/beam_geometry_loads.png')


    # 2. Initialize PyniteFEA Model
    model = FEModel3D()

    # 3. Add Nodes
    # Generate nodes at 100mm (0.1m) segments for detailed results
    node_coords = generate_nodes(span, segment_length=0.1)
    node_names_ordered = list(node_coords.keys()) # N0, n1, n2, ..., N1

    for name, (x, y, z) in node_coords.items():
        model.add_node(name, x, y, z)

    # 4. Define Material and Section (PyniteFEA requires these)
    # Using placeholder values for E and G for timber.
    # These would be more specific if detailed timber design checks were performed.
    # For deflection, E is important. G is for shear deformation (less critical for slender beams).
    # Typical E for C24 might be around 11000 N/mm^2 (11e9 N/m^2).
    # Area = (B*D) mm^2, Iz = (B*D^3)/12 mm^4
    # For Pynite, units should be consistent (e.g., N, m)
    E_val = 11e9  # N/m^2
    G_val = 690e6   # N/m^2 (approx G for timber)
    nu_val = 0.3 # Poisson's ratio (placeholder)
    rho_val = 500 # kg/m^3 (placeholder density for C24)

    model.add_material('TimberC24', E_val, G_val, nu_val, rho_val)

    breadth_m = timber_section['Breadth'] / 1000 * timber_section['No. of Section']
    depth_m = timber_section['Depth'] / 1000
    Area = breadth_m * depth_m
    Iz = (breadth_m * depth_m**3) / 12
    Iy = (depth_m * breadth_m**3) / 12 # Moment of inertia about y-axis
    J = Iy + Iz # Placeholder for torsional constant, less critical for 2D bending

    # PyniteFEA add_section takes name and properties directly
    # For a rectangular section, it primarily needs Area, Iz (bending about z), Iy (bending about y), and J (torsion)
    # The add_rectangle_section was incorrect; instead, we define a general section.
    model.add_section('TimberSection', Area, Iy, Iz, J)


    # 5. Add Members
    # Create one continuous member spanning all nodes for simplicity in PyniteFEA load application
    # Or, create segments between each generated node.
    # For now, one member from N0 to N1, results can be interpolated/queried at internal nodes.
    # Let's try segments for more direct results at n_i nodes.
    for i in range(len(node_names_ordered) - 1):
        start_node = node_names_ordered[i]
        end_node = node_names_ordered[i+1]
        member_name = f"M{i}"
        model.add_member(member_name, start_node, end_node, 'TimberC24', 'TimberSection')


    # 6. Define Support Conditions
    parsed_supports = parse_support_conditions(beam_props['Support conditions'])
    if len(parsed_supports) >= 1:
        model.def_support('N0', parsed_supports[0]['UX'], parsed_supports[0]['UY'], parsed_supports[0]['UZ'],
                                parsed_supports[0]['RX'], parsed_supports[0]['RY'], parsed_supports[0]['RZ'])
    if len(parsed_supports) >= 2:
         model.def_support('N1', parsed_supports[1]['UX'], parsed_supports[1]['UY'], parsed_supports[1]['UZ'],
                                parsed_supports[1]['RX'], parsed_supports[1]['RY'], parsed_supports[1]['RZ'])
    else: # Default for N1 if only one condition provided (e.g. roller)
        model.def_support('N1', False, True, False, False, False, False) # Roller in Y

    # 7. Apply Loads (as separate load cases for Gk, Qk, Sk)
    # PyniteFEA applies loads to members. We have one main member M_main from N0 to N1.
    # Find the main member that spans the whole beam (or the first segment if applying to that)
    # Since we created segments, we need to identify which segment(s) a load applies to.
    # For simplicity of example, assume loads apply to the overall span N0-N1.
    # PyniteFEA's add_member_dist_load and add_member_pt_load refer to the member's local axes.

    # Find the member that starts at N0 to apply loads to, or iterate if loads are complex
    # For UDLs and Point Loads, we need to iterate through members and apply partial loads if needed.

    # Create base load cases
    load_case_map = {} # To store mapping like "Gk1" -> "permanent_Gk1" for Pynite

    # Unfactored Loads for Reactions
    for udl in beam_loads.get('UDL', []):
        lc_name = f"{udl['Load type']}_{udl['Load name']}"
        load_case_map[udl['Load name']] = lc_name

        start_coord = udl['Start'] if udl['Start'] is not None else 0
        end_coord = udl['End'] if udl['End'] is not None else span
        if udl['Full/Partial'].lower() == 'full':
            start_coord = 0
            end_coord = span

        # Apply UDL to relevant segments
        for i in range(len(node_names_ordered) - 1):
            m_name = f"M{i}"
            m_obj = model.members[m_name]
            m_start_x = model.nodes[m_obj.i_node.name].X # Corrected: i_Node -> i_node
            m_end_x = model.nodes[m_obj.j_node.name].X # Corrected: j_Node -> j_node

            # Check if the UDL overlaps with this member segment
            overlap_start = max(start_coord, m_start_x)
            overlap_end = min(end_coord, m_end_x)

            if overlap_end > overlap_start: # There is an overlap
                # PyniteFEA dist load needs start and end magnitude, and positions along member length
                # For a UDL, w_i = w_j = magnitude
                # x_i, x_j are distances from the start of THIS member segment
                load_x_i = max(0, start_coord - m_start_x)
                load_x_j = min(m_obj.L(), end_coord - m_start_x)

                if load_x_j > load_x_i: # Ensure the load portion is valid for this segment
                    # Magnitude is negative for downward loads (local y-axis)
                    # Convert magnitude from kN/m to N/m
                    model.add_member_dist_load(m_name, 'Fy', -abs(udl['Magnitude'])*1000, -abs(udl['Magnitude'])*1000, load_x_i, load_x_j, lc_name)


    for pl in beam_loads.get('Point Load', []):
        lc_name = f"{pl['Load type']}_{pl['Load name']}"
        load_case_map[pl['Load name']] = lc_name
        pl_coord = pl['Start']

        # Apply Point Load to the correct segment
        for i in range(len(node_names_ordered) - 1):
            m_name = f"M{i}"
            m_obj = model.members[m_name] # Corrected: Members -> members
            m_start_x = model.nodes[m_obj.i_node.name].X
            m_end_x = model.nodes[m_obj.j_node.name].X

            if m_start_x <= pl_coord < m_end_x or (np.isclose(pl_coord, m_end_x) and m_end_x == span):
                load_x = pl_coord - m_start_x # Position along the member
                # Magnitude is negative for downward loads
                # Convert magnitude from kN to N
                model.add_member_pt_load(m_name, 'Fy', -abs(pl['Magnitude'])*1000, load_x, lc_name)
                break # Point load applied to one member


    # 8. Define Load Combinations
    # First, create "combinations" for each base load case to ensure they are solved individually
    # This will allow us to get their unfactored reactions.
    for lc_name_in_map, pynite_lc_name in load_case_map.items():
        # Create a new combo name, e.g., " réaction_permanent_Gk1" to avoid clashes if user names combos like base cases
        reaction_combo_name = f"REAC_{pynite_lc_name}"
        model.add_load_combo(reaction_combo_name, {pynite_lc_name: 1.0})

    # ULS Combinations
    for i, combo in enumerate(load_combos['ULS']['Combinations']):
        combo_name = f'ULS{i+1}'
        factors = {}
        # Map load types to their respective load cases generated earlier
        for gk_udl in filter(lambda l: l['Load type'] == 'permanent', beam_loads.get('UDL', [])):
            factors[load_case_map[gk_udl['Load name']]] = combo['Permanent factor']
        for qk_udl in filter(lambda l: l['Load type'] == 'live', beam_loads.get('UDL', [])):
            if qk_udl['Load name'] in load_case_map:
                factors[load_case_map[qk_udl['Load name']]] = combo['Live factor']

        # Point loads - Permanent
        for gk_pl in filter(lambda l: l['Load type'] == 'permanent', beam_loads.get('Point Load', [])):
            if gk_pl['Load name'] in load_case_map:
                factors[load_case_map[gk_pl['Load name']]] = combo['Permanent factor']
        # Point loads - Live
        for qk_pl in filter(lambda l: l['Load type'] == 'live', beam_loads.get('Point Load', [])):
            if qk_pl['Load name'] in load_case_map:
                factors[load_case_map[qk_pl['Load name']]] = combo['Live factor']
        # Point loads - Snow
        for sk_pl in filter(lambda l: l['Load type'] == 'snow', beam_loads.get('Point Load', [])):
            if sk_pl['Load name'] in load_case_map:
                factors[load_case_map[sk_pl['Load name']]] = combo['Snow factor']
        # Add other types (wind) if present in JSON and model (e.g., Point loads - Wind)
        # for wk_pl in filter(lambda l: l['Load type'] == 'wind', beam_loads.get('Point Load', [])):
        #     if wk_pl['Load name'] in load_case_map:
        #         factors[load_case_map[wk_pl['Load name']]] = combo['Wind factor']
        model.add_load_combo(combo_name, factors)

    # SLS Combinations
    for i, combo in enumerate(load_combos['SLS']['Combinations']):
        combo_name = f'SLS{i+1}'
        factors = {}
        # UDLs - Permanent
        for gk_udl in filter(lambda l: l['Load type'] == 'permanent', beam_loads.get('UDL', [])):
            if gk_udl['Load name'] in load_case_map:
                factors[load_case_map[gk_udl['Load name']]] = combo['Permanent factor']
        # UDLs - Live
        for qk_udl in filter(lambda l: l['Load type'] == 'live', beam_loads.get('UDL', [])):
            if qk_udl['Load name'] in load_case_map:
                factors[load_case_map[qk_udl['Load name']]] = combo['Live factor']

        # Point loads - Permanent
        for gk_pl in filter(lambda l: l['Load type'] == 'permanent', beam_loads.get('Point Load', [])):
            if gk_pl['Load name'] in load_case_map:
                factors[load_case_map[gk_pl['Load name']]] = combo['Permanent factor']
        # Point loads - Live
        for qk_pl in filter(lambda l: l['Load type'] == 'live', beam_loads.get('Point Load', [])):
            if qk_pl['Load name'] in load_case_map:
                factors[load_case_map[qk_pl['Load name']]] = combo['Live factor']
        # Point loads - Snow
        for sk_pl in filter(lambda l: l['Load type'] == 'snow', beam_loads.get('Point Load', [])):
            if sk_pl['Load name'] in load_case_map:
                factors[load_case_map[sk_pl['Load name']]] = combo['Snow factor']
        # Point loads - Wind (example)
        # for wk_pl in filter(lambda l: l['Load type'] == 'wind', beam_loads.get('Point Load', [])):
        #     if wk_pl['Load name'] in load_case_map:
        #         factors[load_case_map[wk_pl['Load name']]] = combo['Wind factor']
        model.add_load_combo(combo_name, factors)

    # 9. Analyze
    # For simple beams, check_stability might not be strictly necessary but good for complex models.
    # If it causes issues for this simple case, can be set to False.
    try:
        model.analyze(check_stability=False)
    except Exception as e:
        print(f"Analysis failed: {e}")
        # Fallback if stability check causes issues with simple setup
        # model.analyze(check_stability=False)


    # 10. Extract Results and Plot
    # Node coordinates for plotting
    x_coords = np.array([model.nodes[name].X for name in node_names_ordered]) # Corrected: Nodes -> nodes

    # --- ULS Bending Moment Envelope ---
    uls_combo_names = [f'ULS{i+1}' for i in range(load_combos['ULS']['No. of ULS combinations'])]
    all_uls_moments_Mz = []
    for combo_name in uls_combo_names:
        moments_at_nodes = []
        for node_name in node_names_ordered:
            # Get moment at the start of the member connected TO this node (if j-node)
            # or end of member connected FROM this node (if i-node)
            # PyniteFEA member results are typically at ends. For nodes, need to query connected members.
            # Simpler: get member diagram for each segment and piece together.
            # model.members[m_name].plot_moment('Mz', combo_name) gives plot object.
            # We need data: model.members[m_name].moment_Mz(x, combo_name)

            # For simplicity, let's get moments at the j-end of each segment
            # This means the last node N1 won't have a "start of member" moment from this logic easily
            # Instead, let's query at each node by finding the member ending at it or starting at it.

            current_moment = 0
            if node_name == 'N0': # Start of the beam
                 # Moment at the start of the first member M0
                 if 'M0' in model.members:
                    current_moment = model.members['M0'].moment('Mz', 0, combo_name)  # Corrected method
            elif node_name == 'N1': # End of the beam
                # Moment at the end of the last member
                last_member_name = f"M{len(node_names_ordered)-2}"
                if last_member_name in model.members:
                    current_moment = model.members[last_member_name].moment('Mz', model.members[last_member_name].L(), combo_name) # Corrected method
            else: # Internal nodes
                # Find member segment that *ends* at this node_name (e.g., node n_k is j_Node of M_{k-1})
                # The node index for 'nk' is k. Member M{k-1} connects n{k-1} to nk.
                node_idx_str = node_name[1:] # e.g., "n5" -> "5"
                try:
                    # Find member index that leads to this node
                    # node_names_ordered = [N0, n1, n2, ..., N1]
                    # If node_name is n_idx, it's the j-node of member M_{idx} (if N0 is M0's i-node)
                    # Let's find the actual member index based on node_names_ordered
                    current_node_list_idx = node_names_ordered.index(node_name)
                    member_ending_at_node_name = f"M{current_node_list_idx - 1}"
                    if member_ending_at_node_name in model.members:
                         member_obj = model.members[member_ending_at_node_name]
                         current_moment = member_obj.moment('Mz', member_obj.L(), combo_name) # Corrected method
                    else: # Should not happen if nodes and members align
                        current_moment = 0 # Fallback
                except (ValueError, IndexError):
                     current_moment = 0 # Fallback if node name parsing fails
            moments_at_nodes.append(current_moment / 1000) # Convert N-m to kN-m
        all_uls_moments_Mz.append(moments_at_nodes)

    all_uls_moments_Mz = np.array(all_uls_moments_Mz)
    env_uls_moment_max = np.max(all_uls_moments_Mz, axis=0)
    env_uls_moment_min = np.min(all_uls_moments_Mz, axis=0)
    max_Mz_uls = np.max(env_uls_moment_max)
    min_Mz_uls = np.min(env_uls_moment_min)

    plt.figure(figsize=(10, 6))
    plt.plot(x_coords, env_uls_moment_max, label='Max Envelope Mz (ULS)', color='r')
    plt.plot(x_coords, env_uls_moment_min, label='Min Envelope Mz (ULS)', color='b')
    plt.fill_between(x_coords, env_uls_moment_min, env_uls_moment_max, color='red', alpha=0.1)
    plt.xlabel('Position along beam (m)')
    plt.ylabel('Bending Moment Mz (kN-m)')
    plt.title('ULS Bending Moment Envelope')
    plt.legend()
    plt.grid(True)
    plt.axhline(0, color='black', lw=0.5)
    plt.savefig('plots/bending_moment_uls.png')
    plt.close()

    # --- ULS Shear Force Envelope ---
    all_uls_shears_Fy = []
    for combo_name in uls_combo_names:
        shears_at_nodes_start = [] # Shear at start of segment
        # shears_at_nodes_end = [] # Shear at end of segment
        # Shear can be discontinuous. PyniteFEA gives V at start (x=0) and end (x=L) of member.
        # We will plot two points per segment for shear

        plot_x_shear = []
        plot_y_shear = []

        for i in range(len(node_names_ordered) - 1):
            m_name = f"M{i}"
            member = model.members[m_name]
            x_start_global = model.nodes[member.i_node.name].X # Corrected: i_Node -> i_node
            x_end_global = model.nodes[member.j_node.name].X # Corrected: j_Node -> j_node

            # Shear at start of member segment
            plot_x_shear.append(x_start_global)
            plot_y_shear.append(member.shear('Fy', 0, combo_name) / 1000) # Corrected method; kN
            # Shear at end of member segment
            plot_x_shear.append(x_end_global)
            plot_y_shear.append(member.shear('Fy', member.L(), combo_name) / 1000) # Corrected method; kN
        all_uls_shears_Fy.append(np.array(plot_y_shear)) # Store y-values for envelope calculation

    # This envelope calculation is tricky because x-coords are duplicated for shear steps
    # For simplicity, let's find overall max/min shear from all combos and all points
    # A more robust envelope would require careful handling of the stepped x-coords.

    # Let's plot each ULS shear and then the envelope from those plots' data
    plt.figure(figsize=(10, 6))
    min_shear_overall = float('inf')
    max_shear_overall = float('-inf')

    # Store all (x,y) pairs for shear from all ULS combos
    all_shear_points_for_envelope = []

    for combo_idx, combo_name in enumerate(uls_combo_names):
        current_plot_x_shear = []
        current_plot_y_shear = []
        for i in range(len(node_names_ordered) - 1):
            m_name = f"M{i}"
            member = model.members[m_name]
            x_start_global = model.nodes[member.i_node.name].X # Corrected: i_Node -> i_node
            x_end_global = model.nodes[member.j_node.name].X

            shear_start = member.shear('Fy', 0, combo_name) / 1000 # Corrected method; kN
            shear_end = member.shear('Fy', member.L(), combo_name) / 1000 # Corrected method; kN

            current_plot_x_shear.extend([x_start_global, x_end_global])
            current_plot_y_shear.extend([shear_start, shear_end])

            all_shear_points_for_envelope.append({'x': x_start_global, 'y': shear_start, 'combo': combo_name})
            all_shear_points_for_envelope.append({'x': x_end_global, 'y': shear_end, 'combo': combo_name})

        # plt.plot(current_plot_x_shear, current_plot_y_shear, label=f'{combo_name} Shear Fy', linestyle='--') # Commented out to show only envelope
        min_shear_overall = min(min_shear_overall, min(current_plot_y_shear))
        max_shear_overall = max(max_shear_overall, max(current_plot_y_shear))

    # Create envelope data (simplified: min/max at each unique x-coordinate)
    unique_x_coords_shear = sorted(list(set(p['x'] for p in all_shear_points_for_envelope)))
    env_uls_shear_max_y = []
    env_uls_shear_min_y = []

    for x_val in unique_x_coords_shear:
        y_values_at_x = [p['y'] for p in all_shear_points_for_envelope if np.isclose(p['x'], x_val)]
        env_uls_shear_max_y.append(max(y_values_at_x) if y_values_at_x else 0)
        env_uls_shear_min_y.append(min(y_values_at_x) if y_values_at_x else 0)

    plt.plot(unique_x_coords_shear, env_uls_shear_max_y, color='r', linewidth=1.5, label='Max Envelope Fy (ULS)')
    plt.plot(unique_x_coords_shear, env_uls_shear_min_y, color='b', linewidth=1.5, label='Min Envelope Fy (ULS)')
    plt.fill_between(unique_x_coords_shear, env_uls_shear_min_y, env_uls_shear_max_y, color='blue', alpha=0.1)

    max_V_uls = max(env_uls_shear_max_y) if env_uls_shear_max_y else 0
    min_V_uls = min(env_uls_shear_min_y) if env_uls_shear_min_y else 0

    plt.xlabel('Position along beam (m)')
    plt.ylabel('Shear Force Fy (kN)')
    plt.title('ULS Shear Force Envelope')
    plt.legend()
    plt.grid(True)
    plt.axhline(0, color='black', lw=0.5)
    plt.savefig('plots/shear_force_uls.png')
    plt.close()


    # --- SLS Deflection Envelope (DY) ---
    sls_combo_names = [f'SLS{i+1}' for i in range(load_combos['SLS']['No. of SLS combinations'])]
    all_sls_deflections_dy = []
    for combo_name in sls_combo_names:
        # Get DY (vertical deflection) at each node
        deflections = [model.nodes[name].DY[combo_name] * 1000 for name in node_names_ordered] # Corrected: Nodes -> nodes; Convert m to mm
        all_sls_deflections_dy.append(deflections)

    all_sls_deflections_dy = np.array(all_sls_deflections_dy)
    env_sls_deflection_max = np.max(all_sls_deflections_dy, axis=0) # Max downward deflection (most negative)
    env_sls_deflection_min = np.min(all_sls_deflections_dy, axis=0) # Max upward deflection (most positive or least negative)

    # Typically interested in max magnitude of deflection
    max_abs_deflection_sls = max(np.abs(env_sls_deflection_max).max(), np.abs(env_sls_deflection_min).max())
    # Find the actual max deflection (most negative for downward)
    max_downward_deflection_sls = np.min(env_sls_deflection_min) # Corrected: env_uls_deflection_min -> env_sls_deflection_min


    plt.figure(figsize=(10, 6))
    # Plotting min envelope (max downward deflection) and max envelope (max upward/least downward)
    plt.plot(x_coords, env_sls_deflection_min, label='Min Envelope Dy (Max Downward) (SLS)', color='b')
    plt.plot(x_coords, env_sls_deflection_max, label='Max Envelope Dy (Max Upward) (SLS)', color='r')
    plt.fill_between(x_coords, env_sls_deflection_min, env_sls_deflection_max, color='green', alpha=0.1)
    plt.xlabel('Position along beam (m)')
    plt.ylabel('Deflection Dy (mm)')
    plt.title('SLS Deflection Envelope')
    plt.legend()
    plt.grid(True)
    plt.axhline(0, color='black', lw=0.5)
    plt.gca().invert_yaxis() # Typically deflections are shown positive downwards
    plt.savefig('plots/deflection_sls.png')
    plt.close()

    # --- Unfactored Support Reactions ---
    unfactored_reactions = {}
    # Iterate through the original load case map to get the base PyniteFEA load case names
    # e.g., load_case_map might be {'Gk1': 'permanent_Gk1', 'Qk1': 'live_Qk1'}
    for user_lc_name, pynite_base_lc_name in load_case_map.items():
        reaction_combo_name = f"REAC_{pynite_base_lc_name}" # This is the combo we created for individual analysis

        # Reaction at N0 for this specific "reaction combination"
        rxn_N0_FY = model.nodes['N0'].RxnFY.get(reaction_combo_name, 0) / 1000 # kN
        # Reaction at N1
        rxn_N1_FY = model.nodes['N1'].RxnFY.get(reaction_combo_name, 0) / 1000 # kN

        # Store reactions using the original user-facing load name (e.g., Gk1) or the pynite base lc name for clarity
        unfactored_reactions[user_lc_name] = {'N0_Fy (kN)': rxn_N0_FY, 'N1_Fy (kN)': rxn_N1_FY}


    # 11. Generate DOCX Report
    doc = Document()
    doc.add_heading('Timber Beam Analysis Report', level=0)

    # Beam Geometry
    doc.add_heading('Beam Geometry', level=1)
    doc.add_paragraph(f"Span: {beam_props['Span']} m")
    support_desc = beam_props['Support conditions']
    # Provide a slightly more descriptive interpretation for common cases in the report
    desc_parts = []
    for part in support_desc.split(','):
        if part == "xyz":
            desc_parts.append("Pinned (fixed translation xyz, free rotation)")
        elif part == "y":
            desc_parts.append("Roller (fixed translation y, free translation xz, free rotation)")
        elif part == "xyzXYZ":
            desc_parts.append("Fixed (fixed translation xyz, fixed rotation XYZ)")
        else:
            desc_parts.append(f"Custom ({part})") # Default for other combinations
    doc.add_paragraph(f"Support conditions (N0, N1): {support_desc} (Interpreted as: {', '.join(desc_parts)})")

    # Add Beam Configuration Diagram
    try:
        doc.add_picture('plots/beam_geometry_loads.png', width=Inches(6.5))
    except FileNotFoundError:
        doc.add_paragraph("[Beam geometry and loading diagram not found]")


    # Section Detail
    doc.add_heading('Section Detail', level=1)
    doc.add_paragraph(f"Timber Grade: {timber_section['Grade']}")
    doc.add_paragraph(f"Number of Sections: {timber_section['No. of Section']}")
    doc.add_paragraph(f"Breadth (b): {timber_section['Breadth']} mm (Total: {timber_section['Breadth'] * timber_section['No. of Section']} mm)")
    doc.add_paragraph(f"Depth (h): {timber_section['Depth']} mm")

    # Member details
    doc.add_heading('Member Details', level=1)
    doc.add_paragraph(f"Service Class: {timber_section['Service Class']}")
    doc.add_paragraph(f"Length of Bearing: {timber_section['Length of Bearing']} mm")
    # Load duration is not in the input JSON, so it's omitted or needs a placeholder.
    # For now, omitting. Could say "Load Duration: As per load types (permanent, live, snow)"

    # Applied Loadings
    doc.add_heading('Applied Loadings', level=1)
    # UDLs
    if beam_loads.get('UDL'):
        doc.add_paragraph("Uniformly Distributed Loads (UDLs):")
        udl_table = doc.add_table(rows=1, cols=5)
        udl_table.style = 'Table Grid'
        hdr_cells_udl = udl_table.rows[0].cells
        hdr_cells_udl[0].text = 'Load Name'
        hdr_cells_udl[1].text = 'Load Type'
        hdr_cells_udl[2].text = 'Magnitude (kN/m)'
        hdr_cells_udl[3].text = 'Start (m)'
        hdr_cells_udl[4].text = 'End (m)'
        for udl in beam_loads['UDL']:
            row_cells = udl_table.add_row().cells
            row_cells[0].text = udl['Load name']
            row_cells[1].text = udl['Load type']
            row_cells[2].text = str(udl['Magnitude'])
            row_cells[3].text = str(udl['Start'] if udl['Full/Partial'].lower() == 'partial' else 0)
            row_cells[4].text = str(udl['End'] if udl['Full/Partial'].lower() == 'partial' else beam_props['Span'])
    # Point Loads
    if beam_loads.get('Point Load'):
        doc.add_paragraph("Point Loads:")
        pl_table = doc.add_table(rows=1, cols=4)
        pl_table.style = 'Table Grid'
        hdr_cells_pl = pl_table.rows[0].cells
        hdr_cells_pl[0].text = 'Load Name'
        hdr_cells_pl[1].text = 'Load Type'
        hdr_cells_pl[2].text = 'Magnitude (kN)'
        hdr_cells_pl[3].text = 'Position (m)'
        for pl in beam_loads['Point Load']:
            row_cells = pl_table.add_row().cells
            row_cells[0].text = pl['Load name']
            row_cells[1].text = pl['Load type']
            row_cells[2].text = str(pl['Magnitude'])
            row_cells[3].text = str(pl['Start'])

    # Load Combinations
    doc.add_heading('Load Combinations', level=1)
    add_load_combination_factors(doc, "ULS", load_combos['ULS'])
    add_load_combination_factors(doc, "SLS", load_combos['SLS'])


    # Analysis Results
    doc.add_heading('Analysis Results', level=1)

    doc.add_heading('Maximum Bending Moment (ULS Envelope)', level=2)
    doc.add_paragraph(f"Maximum Positive Bending Moment ( sagging): {max_Mz_uls:.2f} kN-m")
    doc.add_paragraph(f"Maximum Negative Bending Moment (hogging): {min_Mz_uls:.2f} kN-m")
    try:
        doc.add_picture('plots/bending_moment_uls.png', width=Inches(6.0))
    except FileNotFoundError:
        doc.add_paragraph("[Bending moment plot not found]")

    doc.add_heading('Maximum Shear Force (ULS Envelope)', level=2)
    doc.add_paragraph(f"Maximum Positive Shear Force: {max_V_uls:.2f} kN")
    doc.add_paragraph(f"Maximum Negative Shear Force: {min_V_uls:.2f} kN")
    try:
        doc.add_picture('plots/shear_force_uls.png', width=Inches(6.0))
    except FileNotFoundError:
        doc.add_paragraph("[Shear force plot not found]")

    doc.add_heading('Maximum Deflection (SLS Envelope)', level=2)
    # Max downward deflection is the most negative value from env_sls_deflection_min
    doc.add_paragraph(f"Maximum Downward Deflection: {max_downward_deflection_sls:.2f} mm")
    # Could also report max upward if relevant: max_upward_deflection = np.max(env_sls_deflection_max)
    # doc.add_paragraph(f"Maximum Upward Deflection: {max_upward_deflection:.2f} mm")

    try:
        doc.add_picture('plots/deflection_sls.png', width=Inches(6.0))
    except FileNotFoundError:
        doc.add_paragraph("[Deflection plot not found]")

    doc.add_heading('Unfactored Support Reactions (Vertical)', level=2)
    if unfactored_reactions:
        reaction_table = doc.add_table(rows=1, cols=3)
        reaction_table.style = 'Table Grid'
        hdr_cells_rxn = reaction_table.rows[0].cells
        hdr_cells_rxn[0].text = 'Load Case'
        hdr_cells_rxn[1].text = 'Support N0 Reaction (kN)'
        hdr_cells_rxn[2].text = 'Support N1 Reaction (kN)'
        # The keys in unfactored_reactions are now user_lc_name like "Gk1", "Qk1"
        for user_lc_name, rxns in unfactored_reactions.items():
            row_cells = reaction_table.add_row().cells
            row_cells[0].text = user_lc_name # Display the original load name
            row_cells[1].text = f"{rxns['N0_Fy (kN)']:.2f}"
            row_cells[2].text = f"{rxns['N1_Fy (kN)']:.2f}"
    else:
        doc.add_paragraph("No unfactored reactions calculated or available.")

    # Save Document
    doc.save(output_docx)
    print(f"Report saved to {output_docx}")
    print(f"Plots saved in 'plots' directory.")

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description="Run Timber Beam Analysis.")
    parser.add_argument('--input', type=str, default='beam_input.json',
                        help='Path to the input JSON file (default: beam_input.json)')
    parser.add_argument('--output', type=str, default='beam_analysis_report.docx',
                        help='Path to save the output DOCX report (default: beam_analysis_report.docx)')

    args = parser.parse_args()

    run_beam_analysis(input_file=args.input, output_docx=args.output)
    # Example:
    # python beam_analyzer.py --input beam_input_no_pl.json --output report_no_pl.docx
    # python beam_analyzer.py --input beam_input_multi_pl.json --output report_multi_pl.docx
